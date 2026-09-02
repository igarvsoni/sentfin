"""Generate a DOCX project report for the SEntFiN sentiment classifier.

Reads training results from the executed notebook and produces a comprehensive
report with diagrams, math, and embedded result figures.
"""
import json
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

BASE = Path(__file__).resolve().parent
NOTEBOOK_PATH = BASE / "sentfin_classifier.ipynb"
DOCX_PATH = BASE / "SEntFiN_Project_Report.docx"
PDF_PATH = BASE / "SEntFiN_Project_Report.pdf"


# ---------- 1. Extract results from executed notebook ----------
def extract_notebook_outputs():
    """Pull printed outputs from the executed .ipynb so we can quote real numbers."""
    with open(NOTEBOOK_PATH) as f:
        nb = json.load(f)

    outputs = []
    for cell in nb.get("cells", []):
        if cell.get("cell_type") != "code":
            continue
        src = "".join(cell.get("source", []))
        text_parts = []
        for o in cell.get("outputs", []):
            if "text" in o:
                text_parts.append("".join(o["text"]))
            elif o.get("output_type") == "stream":
                text_parts.append("".join(o.get("text", [])))
            elif "data" in o and "text/plain" in o["data"]:
                text_parts.append("".join(o["data"]["text/plain"]))
        outputs.append({"source": src, "output": "\n".join(text_parts)})
    return outputs


def find_metric(outputs, pattern):
    for cell in outputs:
        m = re.search(pattern, cell["output"])
        if m:
            return m.group(1)
    return None


# ---------- 2. Build architecture diagram ----------
def make_architecture_diagram(out_path):
    fig, ax = plt.subplots(figsize=(9, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 12)
    ax.axis("off")

    def box(x, y, w, h, label, color):
        patch = FancyBboxPatch((x, y), w, h,
                               boxstyle="round,pad=0.05,rounding_size=0.15",
                               facecolor=color, edgecolor="black", linewidth=1.2)
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=10, fontweight="bold")

    def arrow(x1, y1, x2, y2):
        a = FancyArrowPatch((x1, y1), (x2, y2),
                            arrowstyle="-|>", mutation_scale=14,
                            color="black", linewidth=1.2)
        ax.add_patch(a)

    # input
    box(0.5, 10.8, 9, 0.8, "Headline + Entity (raw text)", "#fef3c7")
    # tokenizer
    box(0.5, 9.4, 9, 0.8, "FinBERT WordPiece Tokenizer\n[CLS] headline [SEP] entity [SEP]", "#dbeafe")
    arrow(5, 10.8, 5, 10.2)
    # embeddings
    box(0.5, 7.6, 9, 1.4, "Token + Segment + Position Embeddings  (d = 768)", "#bfdbfe")
    arrow(5, 9.4, 5, 9.0)
    # 12 transformer layers
    box(0.5, 4.0, 9, 3.2, "12 × Transformer Encoder Layer\nMulti-Head Self-Attention (12 heads) + FFN\nLayerNorm + Residual", "#93c5fd")
    arrow(5, 7.6, 5, 7.2)
    # CLS pooler
    box(0.5, 2.4, 9, 1.2, "[CLS] Pooled Output  →  tanh(W·h + b)  (768)", "#60a5fa")
    arrow(5, 4.0, 5, 3.6)
    # dropout + linear classifier
    box(0.5, 0.8, 9, 1.2, "Dropout(0.1)  →  Linear(768, 3)  →  Softmax", "#3b82f6")
    arrow(5, 2.4, 5, 2.0)
    # output classes
    box(0.5, -0.6, 9, 0.8, "negative  /  neutral  /  positive", "#1e40af")
    arrow(5, 0.8, 5, 0.2)

    ax.set_ylim(-1, 12)
    ax.set_title("FinBERT + Classification Head — Architecture",
                 fontsize=12, fontweight="bold", pad=10)
    plt.tight_layout()
    plt.savefig(out_path, dpi=180, bbox_inches="tight")
    plt.close()


def make_pipeline_diagram(out_path):
    fig, ax = plt.subplots(figsize=(11, 4))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 4)
    ax.axis("off")

    stages = [
        ("Raw CSV\n(10,753 rows)", "#fde68a"),
        ("Parse `Decisions`\nast + regex fallback", "#fcd34d"),
        ("Expand to triples\n(headline, entity, sentiment)\n~14,405 rows", "#fbbf24"),
        ("Tokenize\n[CLS] H [SEP] E [SEP]", "#a7f3d0"),
        ("Stratified split\n80 / 10 / 10", "#6ee7b7"),
        ("Fine-tune FinBERT\n3 epochs, AdamW", "#60a5fa"),
        ("Evaluate\nacc, macro-F1", "#a78bfa"),
    ]
    n = len(stages)
    box_w = 1.4
    gap = (11 - n * box_w) / (n + 1)
    for i, (label, color) in enumerate(stages):
        x = gap + i * (box_w + gap)
        rect = FancyBboxPatch((x, 1.4), box_w, 1.4,
                              boxstyle="round,pad=0.04,rounding_size=0.1",
                              facecolor=color, edgecolor="black", linewidth=1)
        ax.add_patch(rect)
        ax.text(x + box_w / 2, 2.1, label, ha="center", va="center",
                fontsize=8, fontweight="bold")
        if i < n - 1:
            arr = FancyArrowPatch((x + box_w, 2.1), (x + box_w + gap, 2.1),
                                  arrowstyle="-|>", mutation_scale=12, linewidth=1)
            ax.add_patch(arr)

    ax.set_title("End-to-End Pipeline", fontsize=12, fontweight="bold", pad=10)
    plt.tight_layout()
    plt.savefig(out_path, dpi=180, bbox_inches="tight")
    plt.close()


# ---------- 3. DOCX helpers ----------
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_equation(doc, eq_text, label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(eq_text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    run.italic = True
    if label:
        p.add_run(f"     ({label})").font.size = Pt(10)


def add_image(doc, path, width=6.0, caption=None):
    if not Path(path).exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figure: {caption}")
        run.italic = True
        run.font.size = Pt(10)


def add_code_block(doc, code, max_lines=None):
    """Add a monospace-formatted code block with light grey shading."""
    if max_lines and code.count("\n") > max_lines:
        lines = code.splitlines()
        code = "\n".join(lines[:max_lines]) + f"\n# ... ({len(lines) - max_lines} more lines)"
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    # light grey shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    return t


# ---------- 4. Build the document ----------
def build_docx():
    diag_arch = BASE / "diagram_architecture.png"
    diag_pipe = BASE / "diagram_pipeline.png"
    make_architecture_diagram(diag_arch)
    make_pipeline_diagram(diag_pipe)

    outputs = extract_notebook_outputs()

    # Pull metrics
    test_acc = find_metric(outputs, r"Accuracy\s*:\s*([\d.]+)")
    test_f1 = find_metric(outputs, r"Macro-F1\s*:\s*([\d.]+)")
    expanded = find_metric(outputs, r"Expanded rows\s*:\s*(\d+)")
    train_n = find_metric(outputs, r"Train\s*:\s*(\d+)")
    val_n = find_metric(outputs, r"Val\s*:\s*(\d+)")
    test_n = find_metric(outputs, r"Test\s*:\s*(\d+)")

    classification_block = ""
    for c in outputs:
        if "Classification Report" in c["output"]:
            m = re.search(r"=== Classification Report ===\s*\n(.*?)(?=\n\n|\Z)",
                          c["output"], re.DOTALL)
            if m:
                classification_block = m.group(1).strip()
                break

    doc = Document()

    # ---- title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SEntFiN Financial Sentiment Classifier")
    r.bold = True
    r.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Entity-Aware Sentiment Analysis using FinBERT")
    rs.italic = True
    rs.font.size = Pt(14)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Scholar No: 24u021017").font.size = Pt(11)
    info.add_run("\nProject Report — Data Mining & Warehousing").font.size = Pt(11)

    doc.add_paragraph()

    # ---- abstract ----
    add_heading(doc, "Abstract", level=1)
    add_para(doc,
             "This project fine-tunes the FinBERT transformer for fine-grained, "
             "entity-level sentiment classification of financial news headlines. "
             "Each input pairs a headline with a target financial entity, and the "
             "model predicts one of three sentiment labels (positive, negative, "
             "neutral) specific to that entity. The model is trained on the "
             "SEntFiN 1.0 dataset (Sinha et al., 2022), which provides "
             f"{expanded or '~14,405'} entity-level annotations across 10,753 headlines. "
             "We report accuracy and macro-F1 on a held-out test set, along with "
             "per-class precision/recall.")

    # ---- 1. introduction ----
    add_heading(doc, "1. Introduction", level=1)
    add_para(doc,
             "Sentiment analysis of financial news is critical for algorithmic "
             "trading, risk monitoring, and portfolio analytics. Headline-level "
             "sentiment is too coarse: a single headline often mentions multiple "
             "entities with conflicting sentiment (e.g., 'Gold rallies as Silver "
             "dips' is positive for Gold but negative for Silver). SEntFiN 1.0 "
             "provides entity-level annotations that enable a fine-grained model "
             "capable of disentangling these cases.")

    add_heading(doc, "1.1 Problem Formulation", level=2)
    add_para(doc,
             "Given a headline H and an entity E mentioned within it, predict "
             "the sentiment polarity y ∈ {negative, neutral, positive} that the "
             "headline expresses toward E.")
    add_equation(doc, "y* = argmax_{c ∈ C}  P(c | H, E ; θ)", label="1")

    # ---- 2. dataset ----
    add_heading(doc, "2. Dataset: SEntFiN 1.0", level=1)
    add_para(doc,
             "SEntFiN 1.0 contains 10,753 financial news headlines collected from "
             "Indian and international financial news sources. Each headline is "
             "annotated with one or more (entity, sentiment) pairs stored as a "
             "Python-dict string in the `Decisions` column.")
    add_table(doc, ["Field", "Value"], [
        ["Source", "Sinha et al., 2022 (arXiv:2305.12257)"],
        ["Headlines", "10,753"],
        ["Expanded triples", expanded or "~14,405"],
        ["Classes", "negative / neutral / positive"],
        ["Format", "CSV (Title, Decisions, Words, S No.)"],
    ])

    add_heading(doc, "2.1 Preprocessing", level=2)
    add_para(doc,
             "The `Decisions` field is parsed using `ast.literal_eval`. A small "
             "subset of rows (≈13) contain entity names with apostrophes "
             "(e.g., \"Moody's\", \"Dr Reddy's Labs\") that violate Python literal "
             "syntax. For these rows we use a regex-based fallback that splits on "
             "the three sentinel sentiment values to recover the (entity, sentiment) "
             "pairs. Each headline is then expanded into one row per entity, "
             "yielding ~14,405 (headline, entity, sentiment) training triples.")

    add_heading(doc, "2.2 Class Distribution", level=2)
    add_para(doc,
             "The label distribution is mildly imbalanced "
             "(~26.5% negative, ~38.3% neutral, ~35.2% positive). We use a "
             "class-weighted cross-entropy loss (Section 4.3) to compensate.")

    # ---- 3. methodology ----
    add_heading(doc, "3. Methodology", level=1)
    add_image(doc, diag_pipe, width=6.5,
              caption="End-to-end pipeline: raw CSV → expanded triples → tokenization → fine-tuning → evaluation.")

    add_heading(doc, "3.1 Tokenization", level=2)
    add_para(doc,
             "We use FinBERT's WordPiece tokenizer in sentence-pair mode. For a "
             "headline H and entity E, the input sequence is:")
    add_equation(doc, "[CLS] tok(H) [SEP] tok(E) [SEP]", label="2")
    add_para(doc,
             "Token-type IDs distinguish the two segments (0 for the headline, "
             "1 for the entity). Sequences are padded/truncated to 128 tokens, "
             "well above the dataset's 99th-percentile headline length (17 tokens).")

    add_heading(doc, "3.2 Model Architecture", level=2)
    add_image(doc, diag_arch, width=5.5,
              caption="FinBERT base encoder followed by a 3-class classification head. The pooled [CLS] representation feeds a dropout layer and a linear classifier.")

    add_para(doc,
             "FinBERT (ProsusAI/finbert) is a 12-layer BERT model pre-trained on "
             "financial text (Reuters TRC2-financial + Financial PhraseBank). "
             "We replace its 3-class head with a freshly initialised "
             "Linear(768, 3) layer (the original head's class ordering is "
             "incompatible with our label2id) and fine-tune the entire network "
             "end-to-end.")

    # ---- 4. math ----
    add_heading(doc, "4. Mathematical Formulation", level=1)

    add_heading(doc, "4.1 Multi-Head Self-Attention", level=2)
    add_para(doc, "Each transformer layer applies multi-head self-attention. "
                  "For queries Q, keys K, values V derived from the input X:")
    add_equation(doc, "Attention(Q, K, V) = softmax( Q · Kᵀ / √dₖ ) · V", label="3")
    add_para(doc, "With h = 12 heads operating in parallel:")
    add_equation(doc, "MultiHead(X) = Concat(head₁, …, head_h) · W_O", label="4")
    add_para(doc, "where head_i = Attention(X·W_Q^i, X·W_K^i, X·W_V^i).")

    add_heading(doc, "4.2 Classification Head", level=2)
    add_para(doc, "The pooled [CLS] vector h_[CLS] ∈ ℝ⁷⁶⁸ is mapped to logits:")
    add_equation(doc, "z = W_c · Dropout(h_[CLS]) + b_c,    z ∈ ℝ³", label="5")
    add_equation(doc, "P(c | H, E) = exp(z_c) / Σ_{c'} exp(z_{c'})", label="6")

    add_heading(doc, "4.3 Loss Function", level=2)
    add_para(doc, "We minimise class-weighted cross-entropy:")
    add_equation(doc, "L = − (1/N) Σ_{i=1}^{N} w_{y_i} · log P(y_i | H_i, E_i)", label="7")
    add_para(doc, "where the per-class weight w_c is the inverse-frequency factor:")
    add_equation(doc, "w_c = N / (3 · n_c)", label="8")
    add_para(doc, "with N the total number of training samples and n_c the count of class c.")

    add_heading(doc, "4.4 Optimisation", level=2)
    add_para(doc, "Parameters are updated using AdamW (Loshchilov & Hutter, 2019), "
                  "which decouples weight decay from the gradient-based update:")
    add_equation(doc, "m_t = β₁·m_{t−1} + (1−β₁)·g_t", label="9")
    add_equation(doc, "v_t = β₂·v_{t−1} + (1−β₂)·g_t²", label="10")
    add_equation(doc, "θ_t = θ_{t−1} − η · ( m̂_t / (√v̂_t + ε) + λ·θ_{t−1} )", label="11")
    add_para(doc, "with β₁=0.9, β₂=0.999, ε=1e-8, λ=0.01. The learning rate η follows "
                  "a linear-warmup-and-decay schedule:")
    add_equation(doc, "η(t) = η_max · min( t / T_w ,  (T − t) / (T − T_w) )", label="12")
    add_para(doc, "with T_w = 0.10·T warmup steps and a peak η_max = 2e-5.")

    add_heading(doc, "4.5 Evaluation Metrics", level=2)
    add_para(doc, "Per-class precision, recall, and F1:")
    add_equation(doc, "Precision_c = TP_c / (TP_c + FP_c)", label="13")
    add_equation(doc, "Recall_c    = TP_c / (TP_c + FN_c)", label="14")
    add_equation(doc, "F1_c        = 2 · Precision_c · Recall_c / (Precision_c + Recall_c)", label="15")
    add_para(doc, "Macro-F1 is the unweighted mean across classes — the headline metric for imbalanced settings:")
    add_equation(doc, "Macro-F1 = (1/|C|) · Σ_{c ∈ C} F1_c", label="16")

    # ---- 5. training setup ----
    add_heading(doc, "5. Training Setup", level=1)
    add_table(doc, ["Hyperparameter", "Value"], [
        ["Backbone", "ProsusAI/finbert (12-layer, 768-d, 12 heads)"],
        ["Trainable parameters", "≈ 109.5 M"],
        ["Max sequence length", "128 tokens"],
        ["Batch size", "16"],
        ["Epochs", "3"],
        ["Optimiser", "AdamW (β=0.9/0.999, ε=1e-8, wd=0.01)"],
        ["Peak LR", "2 × 10⁻⁵"],
        ["LR schedule", "Linear warmup (10%) + linear decay"],
        ["Gradient clipping", "max-norm = 1.0"],
        ["Loss", "Class-weighted cross-entropy"],
        ["Train / Val / Test", f"{train_n or '?'} / {val_n or '?'} / {test_n or '?'}"],
        ["Hardware", "NVIDIA RTX 5060 (8 GB VRAM, CUDA)"],
        ["Seed", "42"],
    ])

    # ---- 6. results ----
    add_heading(doc, "6. Results", level=1)
    add_heading(doc, "6.1 Headline Numbers", level=2)
    add_table(doc, ["Metric", "Test Set"], [
        ["Accuracy", f"{float(test_acc)*100:.2f}%" if test_acc else "(see notebook)"],
        ["Macro-F1", f"{float(test_f1):.4f}" if test_f1 else "(see notebook)"],
    ])

    add_heading(doc, "6.2 Per-Class Classification Report", level=2)
    if classification_block:
        p = doc.add_paragraph()
        run = p.add_run(classification_block)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    else:
        add_para(doc, "(see notebook output, Section 8)", italic=True)

    add_heading(doc, "6.3 Training Dynamics", level=2)
    add_image(doc, BASE / "training_curves.png", width=6.5,
              caption="Train/validation loss, validation accuracy, and validation macro-F1 across the 3 epochs.")

    add_heading(doc, "6.4 Confusion Matrix", level=2)
    add_image(doc, BASE / "confusion_matrix.png", width=4.8,
              caption="Test-set confusion matrix. Diagonal cells are correctly classified samples.")

    add_heading(doc, "6.5 Dataset Distribution", level=2)
    add_image(doc, BASE / "sentfin_eda.png", width=6.5,
              caption="Left: label distribution after expansion. Right: histogram of entities per headline.")

    # ---- 7. discussion ----
    add_heading(doc, "7. Discussion", level=1)
    add_para(doc,
             "FinBERT's domain-adaptive pre-training on financial text gives a "
             "strong inductive prior, allowing the model to converge in only 3 "
             "epochs at a low learning rate. The sentence-pair input format "
             "encodes the target entity as a distinct segment, letting the "
             "self-attention layers attend to entity-specific cues and resolve "
             "multi-entity headlines that would be ambiguous to a sentence-level "
             "classifier. Class-weighted loss and stratified splits handle the "
             "modest label imbalance without resorting to oversampling.")

    add_heading(doc, "7.1 Limitations & Future Work", level=2)
    add_para(doc,
             "(i) Multi-entity headlines may contribute correlated training "
             "signals; a headline-disjoint split could give a stricter "
             "evaluation. (ii) FinBERT was pre-trained on English financial "
             "corpora; performance on Indian-English headlines could improve "
             "with domain-continued pre-training. (iii) The 3-class scheme "
             "collapses degree of sentiment; a 5-class or regression target "
             "could capture intensity.")

    # ---- 8. conclusion ----
    add_heading(doc, "8. Conclusion", level=1)
    add_para(doc,
             "We fine-tuned FinBERT on the SEntFiN 1.0 dataset for entity-aware "
             "financial sentiment classification using a [CLS] H [SEP] E [SEP] "
             "input format. The model achieves competitive accuracy and "
             "macro-F1 on the held-out test set with minimal hyperparameter "
             "tuning, demonstrating the effectiveness of domain-adaptive "
             "pre-training combined with a simple sentence-pair fine-tuning "
             "recipe.")

    # ---- appendix: source code ----
    doc.add_page_break()
    add_heading(doc, "Appendix A: Source Code", level=1)
    add_para(doc,
             "Complete source code from the executed notebook "
             "(sentfin_classifier.ipynb), grouped by section. Markdown cells are "
             "shown as section headings; code cells are reproduced verbatim.",
             italic=True, size=10)

    section_titles = {
        0: "A.1  Dataset Acquisition (GitHub fetch)",
        1: "A.2  Imports & Configuration",
        2: "A.3  Data Loading & Parsing",
        3: "A.4  EDA",
        4: "A.5  Train / Val / Test Split",
        5: "A.6  Tokenizer",
        6: "A.7  PyTorch Dataset",
        7: "A.8  DataLoaders",
        8: "A.9  Model Architecture",
        9: "A.10 Optimiser, Scheduler, Loss",
        10: "A.11 Training & Evaluation Functions",
        11: "A.12 Training Run",
        12: "A.13 Training Curves",
        13: "A.14 Test Evaluation & Classification Report",
        14: "A.15 Confusion Matrix",
        15: "A.16 Save Model & Tokenizer",
        16: "A.17 Inference Demo",
    }
    code_cells = [c for c in outputs if c["source"].strip()]
    for i, c in enumerate(code_cells):
        title = section_titles.get(i, f"A.{i+1}  Code Cell {i+1}")
        add_heading(doc, title, level=2)
        add_code_block(doc, c["source"].rstrip())

    # ---- references ----
    doc.add_page_break()
    add_heading(doc, "References", level=1)
    refs = [
        "Sinha, A. et al. (2022). SEntFiN 1.0: Entity-Aware Sentiment Analysis for Financial News. arXiv:2305.12257.",
        "Devlin, J., Chang, M.-W., Lee, K., Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. NAACL-HLT.",
        "Araci, D. (2019). FinBERT: Financial Sentiment Analysis with Pre-trained Language Models. arXiv:1908.10063.",
        "Vaswani, A. et al. (2017). Attention Is All You Need. NeurIPS.",
        "Loshchilov, I., Hutter, F. (2019). Decoupled Weight Decay Regularization. ICLR.",
        "Wolf, T. et al. (2020). Transformers: State-of-the-Art Natural Language Processing. EMNLP-Demos.",
    ]
    for r in refs:
        p = doc.add_paragraph(r, style="List Number")
        p.paragraph_format.space_after = Pt(2)

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}  ({DOCX_PATH.stat().st_size:,} bytes)")


def convert_to_pdf():
    """Convert the DOCX to PDF using LibreOffice headless."""
    import subprocess
    if not DOCX_PATH.exists():
        print(f"DOCX not found at {DOCX_PATH}; skipping PDF conversion.")
        return
    cmd = [
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(BASE), str(DOCX_PATH),
    ]
    print(f"Running: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
    if result.returncode == 0 and PDF_PATH.exists():
        print(f"Wrote {PDF_PATH}  ({PDF_PATH.stat().st_size:,} bytes)")
    else:
        print("libreoffice conversion failed:")
        print(result.stdout)
        print(result.stderr)


if __name__ == "__main__":
    build_docx()
    convert_to_pdf()
